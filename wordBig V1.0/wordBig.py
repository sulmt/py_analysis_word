#!/usr/bin/python3
# -*- coding: utf-8 -*-
# 导入tkinter包，为其定义别名tk
import tkinter as tk
import docx
import xlrd
import time
import sys
import threading
from tkinter import ttk
from docx.shared import Pt
from tkinter import filedialog
from tkinter import messagebox
from docxtpl import DocxTemplate
from PIL import Image, ImageTk



# 定义Application类表示应用/窗口，继承Frame类
class Application(tk.Frame):
    # Application构造函数，master为窗口的父控件
    def __init__(self, master=None):
        # 初始化Application的Frame部分
        tk.Frame.__init__(self, master)
        # 显示窗口，并使用grid布局
        self.grid()
        # 创建控件
        self.create_imput()

    # 创建控件
    def create_imput(self):
        # 创建一个文字为'Quit'，点击会退出的按钮
        # self.quitButton = tk.Button(self, text='Quit', command=self.quit)
        # 显示按钮，并使用grid布局
        # self.quitButton.grid()
        load = Image.open(r'face.jpg')
        render = ImageTk.PhotoImage(load)

        img = tk.Label(image=render)
        img.image = render
        img.grid(row=0, column=0, rowspan=8, columnspan=4)

        self.label1 = tk.Label(text="导入Excel:", fg="black")
        self.label1.grid(row=0, column=4, columnspan=2)

        self.file_import = tk.Button(text="选择文件", command=self.file_picker)
        self.file_import.grid(
            row=0, column=6, columnspan=2, padx=10, sticky='w')

        self.label2 = tk.Label(text="测区名称:", fg="black")
        self.label2.grid(row=2, column=4, columnspan=2)

        self.entry = tk.Entry(width=15)
        self.entry.grid(row=2, column=6, columnspan=2, padx=10, sticky='w')

        self.label3 = tk.Label(text="比例尺:", fg="black")
        self.label3.grid(row=3, column=4, columnspan=2)

        # 创建Radiobutton组
        var = tk.IntVar()
        # 设置组号为1
        var.set(1)
        self.var = var
        self.ratio500 = tk.Radiobutton(
            text="1:500", variable=var, value=1, command=self.select_ratio)
        self.ratio500.grid(row=4, column=4, columnspan=2, padx=5, sticky='w')

        self.ratio1000 = tk.Radiobutton(
            text="1:1000", variable=var, value=2, command=self.select_ratio)
        self.ratio1000.grid(row=4, column=6, columnspan=2, padx=5, sticky='w')

        self.ratio10000 = tk.Radiobutton(
            text="1:10000", variable=var, value=3, command=self.select_ratio)
        self.ratio10000.grid(row=5, column=4, columnspan=2, padx=5, sticky='w')

        self.ratio50000 = tk.Radiobutton(
            text="1:50000", variable=var, value=4, command=self.select_ratio)
        self.ratio50000.grid(row=5, column=6, columnspan=2, padx=5, sticky='w')
        # 比例尺默认选项
        self.scale = "1:500"

        self.btn = tk.Button(text="  执行  ", command=self.path_picker)
        self.btn.grid(row=6, column=4, columnspan=2)

        self.pbar = ttk.Progressbar(
            orient = "horizontal", length = 110, mode = "determinate", value = 0)
        self.pbar.grid(row=6, column=6,columnspan=2)
        self.label4 = tk.Label(
            text="@联系作者：FJSM 汪含秋", fg="black", font=('微软雅黑', 7))
        self.label4.grid(row=7, column=4, columnspan=4)

    # 文件选择
    def file_picker(self):
        filename = filedialog.askopenfilename(filetypes=[("Excel file",
                                                          "*.xls*")])
        if filename != '':
            messagebox.showinfo(title='提示', message="您选择的文件是：" + filename)
            self.excel_path = filename
        else:
            messagebox.showerror(title='错误', message="您未选择任何文件")

    # 路径选择
    def path_picker(self):
        save_path = filedialog.askdirectory()
        if save_path != '':
            messagebox.showinfo(title='提示', message="您选择的路径是：" + save_path)
            self.save_path = save_path
     #       self.analysis_excel(self.excel_path)
            self.pbar.start()
            self.thread1 = threading.Thread(
                target = self.analysis_excel, args=(self.excel_path,))
            self.thread1.setDaemon(True)
            self.thread1.start()
        else:
            messagebox.showerror(title='错误', message="您未选择任何路径")

    # 比例尺选择
    def select_ratio(self):
        selection = self.var.get()
        scale = {
            1: '1:500',
            2: '1:1000',
            3: '1:10000',
            4: '1:50000'
        }
        self.scale = scale.get(selection)
     #   print(self.scale)

    # 检查者+日期 格式转换
    def myfomart(self, head, tail, length):
        blank = ' '
        return head + blank*(length - len(head)*2) + tail
    #汉字 空格推算
    def myfomart2(self, head, length):
        blank = ' '
        return head + blank*(length - len(head)*2)
    #字母数字
    def myfomart3(self, head, length):
        blank = ' '
        return head + blank*(length - len(head))

    # 解析excel表格
    def analysis_excel(self, excel_path):
        # excel
        book = xlrd.open_workbook(excel_path)
        sheet1 = book.sheet_by_index(0)
        sheet_row = sheet1.nrows
        sheet_col = sheet1.ncols
        for r in range(sheet_row):
            if r == 0:
                continue
            data = []
            for c in range(sheet_col):
                value = sheet1.cell_value(r, c)
                if (isinstance(value, float)):
                    value = str(int(value))
                data.append(value)
         #   self.docxtpl_test(data)
            self.thread2 = threading.Thread(
                target=self.docxtpl_test,args=(data,)
            )
            self.thread2.start()
     #   print  ("结束")
        return
     #   messagebox.showinfo(title= '提示', message= '输出完成')



    # 将表格数据填入word
    def docxtpl_test(self, data):
     #   print(data)
     #   print(len(data))
        auth2 = self.myfomart(data[24], data[25], 12)
        check2 = self.myfomart(data[26], data[27], 12)
        auth3 = self.myfomart(data[28], data[29], 12)
        check3 = self.myfomart(data[30], data[31], 12)
        auth4 = self.myfomart(data[32], data[33], 12)
        check4 = self.myfomart(data[34], data[35], 12)
        auth5 = self.myfomart(data[36], data[37], 12)
        check5 = self.myfomart(data[38], data[39], 12)
        method1 = self.myfomart2(data[40], 16)
        soft = self.myfomart3(data[41], 19)
        gs1 = self.myfomart3(data[42], 38)
        method2 = self.myfomart2(data[43], 16)
        gs2 = self.myfomart3(data[44], 18)
        method3 = self.myfomart2(data[45], 16)
        gs3 = self.myfomart3(data[46], 18)
     #   print (data[41])
     #   print (len(data[41]))
     #   print (soft)
     #   print (data[42])
     #   print (len(data[42]))
     #   print (gs1)
     #   print (data[43])
     #   print (method2)
     #   print (len(method2))
     #   print (len(method3))
     #   print (len(gs2))
     #   print (len(gs3))
        doc = DocxTemplate(r'blank_pack.docx')
        context = {
            '图号': data[0],
            '图名': data[1],
            '测区名称': self.entry.get(),
            '比例尺': self.scale,
            '年': data[2],
            'auth0': data[3],
            'check0': data[4],
            'No1': data[5],
            'man1': data[6],
            'chk1': data[7],
            'No2': data[8],
            'man2': data[9],
            'chk2': data[10],
            'No3': data[11],
            'man3': data[12],
            'chk3': data[13],
            'No4': data[14],
            'man4': data[15],
            'chk4': data[16],
            'No5': data[17],
            'man5': data[18],
            'chk5': data[19],
            'No6': data[20],
            'man6': data[21],
            'chk6': data[22],
            'auth1': data[23],
            'auth2': auth2,
            'check2': check2,
            'auth3': auth3,
            'check3': check3,
            'auth4': auth4,
            'check4': check4,
            'auth5': auth5,
            'check5': check5,
            'method1': method1,
            'method2': method2,
            'soft': soft,
            'gs1': gs1,
            'gs2': gs2,
            'method3': method3,
            'gs3': gs3,
        }
        doc.render(context)
        test_path = self.save_path + "/" + data[0] + ".docx"
     #   print (self.save_path)
        doc.save(test_path)
        self.pbar.stop()

        return

def docx_test():
    # excel
    book = xlrd.open_workbook(r'D:\工作\word操作\code_v2\code\data.xlsx')
    sheet1 = book.sheet_by_index(0)
    sheet_row = sheet1.nrows
    sheet_col = sheet1.ncols
    for r in range(sheet_row):
        if r == 0:
            continue
        data = []
        for c in range(sheet_col):
            value = sheet1.cell_value(r, c)
            if (isinstance(value, float)):
                value = str(int(value))
            data.append(value)
    #print(data)
    # word
    document = docx.Document(r'D:/code/py_code/pyGUI/temple.docx')
    tables = document.tables
    homeTable = tables[0]
    # 图号
    run = homeTable.cell(0, 0).paragraphs[0].add_run(data[0])
    run.font.name = u'微软雅黑'
    run.font.size = Pt(15)
    # 图名
    run = homeTable.cell(1, 0).paragraphs[0].add_run(data[1])
    run.font.name = u'微软雅黑'
    run.font.size = Pt(15)
    # 测区名称
    run = homeTable.cell(2, 0).paragraphs[0].add_run(data[2])
    run.font.name = u'微软雅黑'
    run.font.size = Pt(15)
    # 比例尺

    # 年
    run = tables[1].cell(0, 0).paragraphs[0].add_run(data[3])
    run.font.name = u'微软雅黑'
    run.font.size = Pt(15)

    annotateTable = tables[3]
    # 循环填表
    # for i in range(6):
    #     # 绘图编号1
    run = annotateTable.cell(1, 1).paragraphs[0].add_run(data[6])
    run.font.name = u'宋体'
    run.font.size = Pt(12)
    #     # 作业者1
    #     run = annotateTable.cell(1, 2).paragraphs[0].add_run(data[7 + i*3])
    #     run.font.name = u'宋体'
    #     run.font.size = Pt(12)
    #     # 检查者1
    #     run = annotateTable.cell(1, 3).paragraphs[0].add_run(data[8 + i*3])
    #     run.font.name = u'宋体'
    #     run.font.size = Pt(12)
    paragraphs = document.paragraphs
    i = 0
    t = 0
    blank = ' '
    for p in paragraphs:
     #   print(i)
        i = i + 1
        if ('作业方法' in p.text):
            p.text = ''
            run = p.add_run('作业方法：')
            run.font.name = u'宋体'
            run.font.size = Pt(12)
            workway = '随机数测试'
            workway = workway + blank * (16 - len(workway) * 2)
            run = p.add_run(workway)
            run.font.name = u'宋体'
            run.font.size = Pt(12)
            run.font.underline = True
            run = p.add_run(' 质量检查方法（软件）：')
            run.font.name = u'宋体'
            run.font.size = Pt(12)
            checkway = '测试随机长度'
            checkway = checkway + blank * (21 - len(checkway) * 2)
            run = p.add_run(checkway)
            run.font.name = u'宋体'
            run.font.size = Pt(12)
            run.font.underline = True
        if ('填表者' in p.text):
            if (t == 0 or t == 2):
                # p.text = ''
                run = p.add_run(data[4])
                run.font.name = u'宋体'
                run.font.size = Pt(12)
            if (t == 1):
                pass
        if ('检查者' in p.text):
            if (t == 0 or t == 2):
                run = p.add_run(data[5])
                run.font.name = u'宋体'
                run.font.size = Pt(12)
    timestamp = time.strftime('%Y%m%d%H%M%S', time.localtime(time.time()))
    test_path = "D:\工作\word操作\code_v2\code" + timestamp + ".docx"
    document.save(test_path)

# docx_test()

# 创建一个Application对象app
app = Application()
# 设置窗口标题为'First Tkinter'
app.master.title = 'First Tkinter'
# 主循环开始
app.mainloop()
#messagebox.showinfo("ok")
